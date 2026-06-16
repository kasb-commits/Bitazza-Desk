"""
Knowledge Base management routes.

Supports ingesting content from:
- URLs (scraped via BeautifulSoup)
- PDF files (via pypdf)
- DOCX files (via python-docx)

All content is chunked and stored in ChromaDB so the RAG retriever
automatically picks it up when answering customer queries.
"""
import io
import logging
import os
import uuid

import requests as _requests
from bs4 import BeautifulSoup
from fastapi import APIRouter, Depends, Form, HTTPException, Request, UploadFile, File
from pydantic import BaseModel

from fastapi import BackgroundTasks
from api.middleware.auth import get_user_id, get_optional_user_id
from db.conversation_store import (
    create_knowledge_item,
    list_knowledge_items,
    get_knowledge_item,
    delete_knowledge_item,
    update_knowledge_item_citations,
    update_knowledge_item_citations_manual,
    get_all_knowledge_source_refs,
    create_knowledge_item_version,
    activate_knowledge_version,
    fail_knowledge_version,
    get_knowledge_item_family,
    delete_knowledge_item_chain,
)
from db.vector_store import upsert_documents, delete_by_metadata, get_chunks_by_item
from engine.kb_classifier import classify_kb_item

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/knowledge", tags=["knowledge"])


# ── Text chunking ─────────────────────────────────────────────────────────────

def _chunk_text(text: str, size: int = 800, overlap: int = 100) -> list[str]:
    """Split text into overlapping chunks of ~size characters."""
    # Prefer paragraph splits first
    paragraphs: list[str] = []
    for para in text.split("\n\n"):
        para = para.strip()
        if para:
            paragraphs.append(para)

    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= size:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            # If a single paragraph is > size, split by words
            if len(para) > size:
                words = para.split()
                buf = ""
                for w in words:
                    if len(buf) + len(w) + 1 <= size:
                        buf = (buf + " " + w).strip()
                    else:
                        if buf:
                            chunks.append(buf)
                        # Carry overlap
                        overlap_words = buf.split()[-max(1, overlap // 6):]
                        buf = " ".join(overlap_words) + " " + w
                if buf:
                    current = buf
            else:
                # Carry overlap from previous chunk
                if chunks:
                    prev_words = chunks[-1].split()
                    carry = " ".join(prev_words[-max(1, overlap // 6):])
                    current = (carry + "\n\n" + para).strip()
                else:
                    current = para

    if current:
        chunks.append(current)

    return [c for c in chunks if c.strip()]


# ── URL ingestion ─────────────────────────────────────────────────────────────

class AddUrlRequest(BaseModel):
    url: str


def _classify_and_persist(item_id: int, title: str, chunks: list[str]) -> None:
    """Background task: classify a KB item and persist citation metadata."""
    try:
        result = classify_kb_item(item_id, title, chunks)
        update_knowledge_item_citations(
            item_id,
            result["categories"],
            result["keywords"],
            result["coverage_score"],
        )
        logger.info("KB classifier: item %d classified — %s", item_id, result["categories"])
    except Exception:
        logger.exception("KB classifier: failed for item %d", item_id)


@router.post("/url")
def add_url(
    body: AddUrlRequest,
    background_tasks: BackgroundTasks,
    agent_id: str = Depends(get_user_id),
):
    """Scrape a URL, chunk its content, and store in the vector DB."""
    url = body.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=422, detail="URL must start with http:// or https://")

    try:
        resp = _requests.get(url, timeout=15, headers={"User-Agent": "CSBot/1.0"})
        resp.raise_for_status()
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to fetch URL: {exc}") from exc

    soup = BeautifulSoup(resp.text, "html.parser")

    # Remove boilerplate elements
    for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # Title
    title_tag = soup.find("title") or soup.find("h1")
    title = title_tag.get_text(strip=True) if title_tag else url

    body_text = soup.get_text(separator="\n", strip=True)
    if not body_text.strip():
        raise HTTPException(status_code=422, detail="No readable content found at URL")

    chunks = _chunk_text(body_text)
    if not chunks:
        raise HTTPException(status_code=422, detail="Could not extract any text from URL")

    # Persist to DB first to get the item ID
    item = create_knowledge_item(
        title=title[:255],
        source_type="url",
        source_ref=url,
        chunk_count=len(chunks),
        created_by=agent_id,
    )
    item_id = str(item["id"])

    # Upsert chunks to vector store
    docs = [
        {
            "id": f"kb_{item_id}_{i}",
            "text": chunk,
            "metadata": {
                "knowledge_item_id": item_id,
                "source": url,
                "source_type": "url",
                "title": title[:255],
                "chunk_index": i,
                "status": "ACTIVE",
            },
        }
        for i, chunk in enumerate(chunks)
    ]
    try:
        upsert_documents(docs)
    except Exception as exc:
        logger.exception("Vector upsert failed for knowledge item %s", item_id)
        raise HTTPException(status_code=500, detail=f"Vector store error: {exc}") from exc

    background_tasks.add_task(_classify_and_persist, item["id"], title, chunks)
    return item


# ── File upload ingestion ─────────────────────────────────────────────────────

@router.post("/upload")
async def upload_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    agent_id: str = Depends(get_user_id),
):
    """Upload a PDF or DOCX file, extract its text, and store in the vector DB."""
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    content_type = (file.content_type or "").lower()

    is_pdf = ext == "pdf" or "pdf" in content_type
    is_docx = ext == "docx" or "wordprocessingml" in content_type

    if not (is_pdf or is_docx):
        raise HTTPException(
            status_code=422,
            detail="Only PDF (.pdf) and Word (.docx) files are supported"
        )

    raw = await file.read()
    buf = io.BytesIO(raw)
    text = ""

    if is_pdf:
        try:
            import pypdf  # type: ignore
        except ImportError:
            raise HTTPException(
                status_code=422,
                detail="PDF support requires pypdf. Run: pip install pypdf"
            )
        try:
            reader = pypdf.PdfReader(buf)
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(p for p in pages if p.strip())
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not read PDF: {exc}") from exc

    elif is_docx:
        try:
            import docx  # type: ignore  # python-docx
        except ImportError:
            raise HTTPException(
                status_code=422,
                detail="DOCX support requires python-docx. Run: pip install python-docx"
            )
        try:
            doc = docx.Document(buf)
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not read DOCX: {exc}") from exc

    if not text.strip():
        raise HTTPException(status_code=422, detail="No readable text found in the uploaded file")

    chunks = _chunk_text(text)
    if not chunks:
        raise HTTPException(status_code=422, detail="Could not extract any text from the file")

    title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ") if filename else "Uploaded document"

    item = create_knowledge_item(
        title=title[:255],
        source_type=ext if ext in ("pdf", "docx") else "pdf",
        source_ref=filename,
        chunk_count=len(chunks),
        created_by=agent_id,
    )
    item_id = str(item["id"])

    docs = [
        {
            "id": f"kb_{item_id}_{i}",
            "text": chunk,
            "metadata": {
                "knowledge_item_id": item_id,
                "source": filename,
                "source_type": ext,
                "title": title[:255],
                "chunk_index": i,
                "status": "ACTIVE",
            },
        }
        for i, chunk in enumerate(chunks)
    ]
    try:
        upsert_documents(docs)
    except Exception as exc:
        logger.exception("Vector upsert failed for knowledge item %s", item_id)
        raise HTTPException(status_code=500, detail=f"Vector store error: {exc}") from exc

    background_tasks.add_task(_classify_and_persist, item["id"], title, chunks)
    return item


# ── List ──────────────────────────────────────────────────────────────────────

@router.get("")
def get_knowledge_items(_agent_id: str = Depends(get_user_id)):
    """List all knowledge base items."""
    items = list_knowledge_items()
    return {"items": items}


# ── Preview chunks ────────────────────────────────────────────────────────────

@router.get("/{item_id}/chunks")
def get_chunks(item_id: int, _agent_id: str = Depends(get_optional_user_id)):
    """Return all indexed text chunks for a knowledge item."""
    item = get_knowledge_item(item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Knowledge item not found")
    try:
        chunks = get_chunks_by_item(item_id)
        return {"item_id": item_id, "chunks": chunks}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


# ── Bulk blog crawl ───────────────────────────────────────────────────────────

class BulkBlogRequest(BaseModel):
    index_url: str
    base_url: str = ""          # inferred from index_url if empty
    article_path_prefix: str = "/blog/"  # only follow links containing this

@router.post("/admin/bulk-ingest-blog")
def bulk_ingest_blog(body: BulkBlogRequest, request: Request):
    """
    Crawl a blog index page, find all article links, and ingest each one.
    Protected by INTERNAL_SERVICE_TOKEN.
    Returns a summary of what was ingested / skipped / failed.
    """
    from config import settings as _cfg
    import time as _time

    auth = request.headers.get("Authorization", "")
    token = auth.removeprefix("Bearer ").strip()
    if not _cfg.INTERNAL_SERVICE_TOKEN or token != _cfg.INTERNAL_SERVICE_TOKEN:
        raise HTTPException(status_code=403, detail="Forbidden")

    index_url = body.index_url.strip()
    base_url = body.base_url.strip() or "/".join(index_url.split("/")[:3])
    prefix = body.article_path_prefix

    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
    }

    # Fetch index page
    try:
        resp = _requests.get(index_url, timeout=20, headers=headers)
        resp.raise_for_status()
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to fetch index: {exc}")

    soup = BeautifulSoup(resp.text, "html.parser")

    # Collect all article links
    article_urls: set[str] = set()
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if prefix in href:
            if href.startswith("http"):
                full = href
            elif href.startswith("/"):
                full = base_url + href
            else:
                continue
            # Skip the index page itself
            if full.rstrip("/") == index_url.rstrip("/"):
                continue
            article_urls.add(full.split("?")[0].split("#")[0])

    if not article_urls:
        return {"error": "No article links found — the page structure may require JavaScript rendering", "index_url": index_url}

    # Get already-indexed URLs to skip duplicates — must check all statuses (including ARCHIVED)
    # so that archived blog URLs from previous override cycles are not re-ingested.
    existing = get_all_knowledge_source_refs()

    results = {"ingested": [], "skipped_duplicate": [], "failed": []}

    for url in sorted(article_urls):
        if url in existing:
            results["skipped_duplicate"].append(url)
            continue

        try:
            page = _requests.get(url, timeout=20, headers=headers)
            page.raise_for_status()
        except Exception as exc:
            logger.warning("bulk-ingest-blog: fetch failed %s — %s", url, exc)
            results["failed"].append({"url": url, "reason": str(exc)})
            continue

        page_soup = BeautifulSoup(page.text, "html.parser")
        for tag in page_soup.find_all(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        title_tag = page_soup.find("title") or page_soup.find("h1")
        title = title_tag.get_text(strip=True) if title_tag else url
        body_text = page_soup.get_text(separator="\n", strip=True)

        if not body_text.strip():
            results["failed"].append({"url": url, "reason": "no readable content"})
            continue

        chunks = _chunk_text(body_text)
        if not chunks:
            results["failed"].append({"url": url, "reason": "chunking produced nothing"})
            continue

        item = create_knowledge_item(
            title=title[:255],
            source_type="url",
            source_ref=url,
            chunk_count=len(chunks),
            created_by=None,
        )
        item_id = str(item["id"])

        docs = [
            {
                "id": f"kb_{item_id}_{i}",
                "text": chunk,
                "metadata": {
                    "knowledge_item_id": item_id,
                    "source": url,
                    "source_type": "url",
                    "title": title[:255],
                    "chunk_index": i,
                    "status": "ACTIVE",
                },
            }
            for i, chunk in enumerate(chunks)
        ]
        try:
            upsert_documents(docs)
            results["ingested"].append({"url": url, "title": title[:80], "chunks": len(chunks)})
            logger.info("bulk-ingest-blog: ingested %s (%d chunks)", url, len(chunks))
        except Exception as exc:
            logger.exception("bulk-ingest-blog: upsert failed %s", url)
            results["failed"].append({"url": url, "reason": str(exc)})

        _time.sleep(0.5)

    return {
        "summary": {
            "found": len(article_urls),
            "ingested": len(results["ingested"]),
            "skipped_duplicate": len(results["skipped_duplicate"]),
            "failed": len(results["failed"]),
        },
        "ingested": results["ingested"],
        "failed": results["failed"],
        "skipped_duplicate": results["skipped_duplicate"],
    }


# ── Migrate ChromaDB → pgvector (one-off admin) ───────────────────────────────

@router.post("/admin/migrate-from-chroma")
def migrate_from_chroma(request: Request):
    """
    One-off endpoint: reads all docs from ChromaDB and upserts into vector_embeddings.
    Protected by INTERNAL_SERVICE_TOKEN (Authorization: Bearer <token>).
    Safe to call multiple times — uses ON CONFLICT DO UPDATE.
    """
    from config import settings as _cfg

    # Auth: require INTERNAL_SERVICE_TOKEN
    auth = request.headers.get("Authorization", "")
    token = auth.removeprefix("Bearer ").strip()
    if not _cfg.INTERNAL_SERVICE_TOKEN or token != _cfg.INTERNAL_SERVICE_TOKEN:
        raise HTTPException(status_code=403, detail="Forbidden")

    chroma_path = _cfg.CHROMA_PATH if hasattr(_cfg, "CHROMA_PATH") else os.environ.get("CHROMA_PATH", "./data/chroma")

    if not os.path.isdir(chroma_path):
        raise HTTPException(
            status_code=422,
            detail=f"ChromaDB path not found: {chroma_path!r} — volume may be empty or path wrong",
        )

    try:
        import chromadb as _chromadb
    except ImportError:
        raise HTTPException(status_code=500, detail="chromadb package not installed on this server")

    try:
        chroma_client = _chromadb.PersistentClient(path=chroma_path)
        col = chroma_client.get_collection("knowledge_base")
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not open ChromaDB collection: {exc}")

    total = col.count()
    if total == 0:
        return {"migrated": 0, "skipped": 0, "message": "ChromaDB collection is empty — no data to recover"}

    # Page through ChromaDB
    PAGE = 500
    all_docs: list[dict] = []
    offset = 0
    while True:
        result = col.get(limit=PAGE, offset=offset, include=["documents", "metadatas"])
        ids       = result.get("ids", [])
        documents = result.get("documents", [])
        metadatas = result.get("metadatas", [])
        if not ids:
            break
        for chroma_id, text, meta in zip(ids, documents, metadatas):
            if text and text.strip():
                all_docs.append({"id": chroma_id, "text": text, "metadata": meta or {}})
        offset += len(ids)
        if len(ids) < PAGE:
            break

    skipped = total - len(all_docs)

    # Upsert in batches of 20 (Gemini embedding batch limit)
    BATCH = 20
    migrated = 0
    errors: list[str] = []
    for start in range(0, len(all_docs), BATCH):
        batch = all_docs[start : start + BATCH]
        try:
            upsert_documents(batch)
            migrated += len(batch)
        except Exception as exc:
            logger.exception("migrate-from-chroma: batch %d–%d failed", start, start + BATCH)
            errors.append(str(exc))

    logger.info("migrate-from-chroma: %d migrated, %d skipped empty, %d errors", migrated, skipped, len(errors))
    return {
        "chroma_total": total,
        "migrated": migrated,
        "skipped_empty": skipped,
        "errors": errors[:10],
    }


# ── Delete ────────────────────────────────────────────────────────────────────

@router.delete("/{item_id}", status_code=204)
def delete_item(item_id: int, _agent_id: str = Depends(get_user_id)):
    """Delete a knowledge item and its entire version chain (all versions + all chunks)."""
    item = get_knowledge_item(item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Knowledge item not found")

    if item.get("status") == "ARCHIVED":
        raise HTTPException(status_code=409, detail="Cannot delete an archived version directly. Delete the active version to remove the entire chain.")

    root_id = item.get("parent_id") or item_id
    deleted_ids = delete_knowledge_item_chain(root_id)
    logger.info("Deleted knowledge item chain root=%d — %d versions removed", root_id, len(deleted_ids))


# ── Citations PATCH ───────────────────────────────────────────────────────────

class CitationsPatchRequest(BaseModel):
    categories: list[str]
    keywords: list[str]


@router.patch("/{item_id}/citations")
def patch_citations(
    item_id: int,
    body: CitationsPatchRequest,
    agent_id: str = Depends(get_user_id),
):
    """Manually override AI-generated citations for a KB item. Locks the item against reclassification."""
    item = get_knowledge_item(item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Knowledge item not found")
    if item.get("status") != "ACTIVE":
        raise HTTPException(status_code=409, detail="Cannot edit citations on a non-active version")

    update_knowledge_item_citations_manual(item_id, body.categories, body.keywords, agent_id)
    updated = get_knowledge_item(item_id)
    return updated


# ── Version history ───────────────────────────────────────────────────────────

@router.get("/{item_id}/versions")
def get_versions(item_id: int, _agent_id: str = Depends(get_user_id)):
    """Return the full version history for a KB document family."""
    item = get_knowledge_item(item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Knowledge item not found")
    root_id = item.get("parent_id") or item_id
    versions = get_knowledge_item_family(root_id)
    return {"versions": versions}


# ── Version override ──────────────────────────────────────────────────────────

@router.post("/{item_id}/override")
async def override_knowledge_item(
    item_id: int,
    background_tasks: BackgroundTasks,
    url: str = Form(default=""),
    change_notes: str = Form(default=""),
    file: UploadFile | None = File(default=None),
    agent_id: str = Depends(get_user_id),
):
    """
    Replace a KB item with a new version (URL or file).
    The old version is atomically archived only after the new one is successfully ingested.
    If any step fails after the new DB row is created, it is marked FAILED and the old
    version stays ACTIVE.
    """
    # Step 1 — load and validate the current active item
    old_item = get_knowledge_item(item_id)
    if not old_item:
        raise HTTPException(status_code=404, detail="Knowledge item not found")
    if old_item.get("status") != "ACTIVE":
        raise HTTPException(status_code=409, detail="Only ACTIVE items can be overridden")

    # Step 2 — resolve root_id for the version chain
    root_id: int = old_item.get("parent_id") or old_item["id"]

    # Step 3 — compute next version number
    family = get_knowledge_item_family(root_id)
    next_version = max((v["version_number"] for v in family), default=1) + 1

    # Step 4 — create new DB row with status=PROCESSING
    new_item = create_knowledge_item_version(
        root_id=root_id,
        version_number=next_version,
        title=old_item["title"],
        source_type=old_item["source_type"],
        source_ref=old_item.get("source_ref"),
        chunk_count=0,
        created_by=agent_id,
        change_notes=change_notes,
    )
    new_id: int = new_item["id"]

    try:
        # Step 5 — ingest content (URL or file)
        if file and file.filename:
            filename = file.filename or ""
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            content_type = (file.content_type or "").lower()
            is_pdf = ext == "pdf" or "pdf" in content_type
            is_docx = ext == "docx" or "wordprocessingml" in content_type
            if not (is_pdf or is_docx):
                raise HTTPException(status_code=422, detail="Only PDF and DOCX files are supported")

            raw = await file.read()
            buf = io.BytesIO(raw)
            text = ""
            if is_pdf:
                try:
                    import pypdf
                    reader = pypdf.PdfReader(buf)
                    text = "\n\n".join(p.extract_text() or "" for p in reader.pages if (p.extract_text() or "").strip())
                except Exception as exc:
                    raise HTTPException(status_code=422, detail=f"Could not read PDF: {exc}") from exc
            elif is_docx:
                try:
                    import docx
                    doc = docx.Document(buf)
                    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception as exc:
                    raise HTTPException(status_code=422, detail=f"Could not read DOCX: {exc}") from exc
            if not text.strip():
                raise HTTPException(status_code=422, detail="No readable text found in the uploaded file")
            source_ref = filename
            title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ")
        elif url:
            url = url.strip()
            if not url.startswith(("http://", "https://")):
                raise HTTPException(status_code=422, detail="URL must start with http:// or https://")
            try:
                resp = _requests.get(url, timeout=15, headers={"User-Agent": "CSBot/1.0"})
                resp.raise_for_status()
            except Exception as exc:
                raise HTTPException(status_code=422, detail=f"Failed to fetch URL: {exc}") from exc
            soup = BeautifulSoup(resp.text, "html.parser")
            for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()
            title_tag = soup.find("title") or soup.find("h1")
            title = title_tag.get_text(strip=True) if title_tag else url
            text = soup.get_text(separator="\n", strip=True)
            if not text.strip():
                raise HTTPException(status_code=422, detail="No readable content found at URL")
            source_ref = url
        else:
            raise HTTPException(status_code=422, detail="Provide either a file upload or a url field")

        chunks = _chunk_text(text)
        if not chunks:
            raise HTTPException(status_code=422, detail="Could not extract any text from the provided content")

        # Step 6 — upsert new chunks with status=ACTIVE in metadata
        new_id_str = str(new_id)
        docs = [
            {
                "id": f"kb_{new_id_str}_{i}",
                "text": chunk,
                "metadata": {
                    "knowledge_item_id": new_id_str,
                    "source": source_ref,
                    "source_type": old_item["source_type"],
                    "title": title[:255],
                    "chunk_index": i,
                    "status": "ACTIVE",
                },
            }
            for i, chunk in enumerate(chunks)
        ]
        upsert_documents(docs)

        # Step 7 — atomically archive old, activate new, mark old chunks ARCHIVED
        activate_knowledge_version(new_id, item_id)

    except HTTPException:
        fail_knowledge_version(new_id)
        raise
    except Exception as exc:
        fail_knowledge_version(new_id)
        logger.exception("Override failed for item %d (new version %d)", item_id, new_id)
        raise HTTPException(status_code=500, detail=f"Override failed: {exc}") from exc

    # Step 8 — schedule citation classification for new version
    background_tasks.add_task(_classify_and_persist, new_id, title[:255], chunks)

    # Step 9 — return the new active item
    return get_knowledge_item(new_id)


# ── Admin: classify-all ───────────────────────────────────────────────────────

@router.post("/admin/classify-all")
def admin_classify_all(_agent_id: str = Depends(get_user_id)):
    """
    Backfill citation metadata for all KB items that have no categories yet.
    Processes: citations_source='pending' OR (citations_source='ai' with empty categories).
    Skips: manually edited items (citations_source='manual').
    """

    import psycopg2
    import psycopg2.extras
    from config import settings as _cfg
    from db.vector_store import get_chunks_by_item as _get_chunks

    try:
        conn = psycopg2.connect(_cfg.DATABASE_URL, cursor_factory=psycopg2.extras.RealDictCursor)
        cur = conn.cursor()
        cur.execute(
            """SELECT id, title FROM knowledge_items
               WHERE status = 'ACTIVE'
                 AND citations_source != 'manual'
                 AND (citations_source = 'pending'
                      OR (citations_source = 'ai' AND (citation_categories IS NULL OR citation_categories = '{}')))"""
        )
        pending = cur.fetchall()
        conn.close()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DB error: {exc}") from exc

    classified = 0
    failed = 0
    for row in pending:
        try:
            chunks_data = _get_chunks(row["id"])
            chunks = [c["text"] for c in chunks_data]
            result = classify_kb_item(row["id"], row["title"], chunks)
            update_knowledge_item_citations(row["id"], result["categories"], result["keywords"], result["coverage_score"])
            classified += 1
        except Exception:
            logger.exception("admin_classify_all: failed for item %d", row["id"])
            failed += 1

    return {"total_pending": len(pending), "classified": classified, "failed": failed}
